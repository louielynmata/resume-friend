import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from backend.services.document_service import (
    _build_cover_letter_docx,
    _build_resume_docx,
)


RESUME_TEXT = """NAME: Alex Example
ROLE: Product Designer
CONTACT: alex@example.com

WORK EXPERIENCE
PRODUCT DESIGNER
Example Studio | Calgary
2020 - Present
● Built accessible customer workflows.

EDUCATION
Example University | Design Diploma
Second University | Arts Degree

CERTIFICATIONS
● Example Certification
"""

COVER_LETTER_TEXT = """Cover Letter

Dear Hiring Team,

I design accessible products and would bring that focus to this role.

I would welcome the opportunity to contribute to the team.

Sincerely,
Alex Example
"""


def paragraph_starting_with(document: Document, prefix: str):
    return next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(prefix)
    )


class DocumentLayoutTests(unittest.TestCase):
    def test_resume_keeps_entry_headers_with_first_content(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "resume.docx"
            _build_resume_docx(RESUME_TEXT, path)
            document = Document(path)

            work_header = paragraph_starting_with(document, "WORK EXPERIENCE")
            role_header = paragraph_starting_with(document, "PRODUCT DESIGNER")
            company = paragraph_starting_with(document, "Example Studio")
            date_range = paragraph_starting_with(document, "2020 - Present")
            first_bullet = paragraph_starting_with(
                document,
                "Built accessible customer workflows.",
            )

            self.assertTrue(work_header.paragraph_format.keep_with_next)
            self.assertTrue(role_header.paragraph_format.keep_with_next)
            self.assertTrue(company.paragraph_format.keep_with_next)
            self.assertTrue(date_range.paragraph_format.keep_with_next)
            self.assertTrue(first_bullet.paragraph_format.keep_together)

    def test_resume_keeps_education_group_but_not_next_section(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "resume.docx"
            _build_resume_docx(RESUME_TEXT, path)
            document = Document(path)

            education = paragraph_starting_with(document, "EDUCATION")
            first_school = paragraph_starting_with(document, "Example University")
            second_school = paragraph_starting_with(document, "Second University")

            self.assertTrue(education.paragraph_format.keep_with_next)
            self.assertTrue(first_school.paragraph_format.keep_with_next)
            self.assertFalse(second_school.paragraph_format.keep_with_next)
            self.assertEqual(education.paragraph_format.space_before.pt, 4)
            self.assertEqual(education.paragraph_format.space_after.pt, 1)

    def test_cover_letter_uses_compact_spacing_and_keeps_closing_together(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "cover-letter.docx"
            _build_cover_letter_docx(COVER_LETTER_TEXT, path)
            document = Document(path)

            heading = paragraph_starting_with(document, "Cover Letter")
            greeting = paragraph_starting_with(document, "Dear Hiring Team")
            opening = paragraph_starting_with(document, "I design accessible")
            closing = paragraph_starting_with(document, "Sincerely,")
            name = paragraph_starting_with(document, "Alex Example")

            self.assertTrue(heading.paragraph_format.keep_with_next)
            self.assertTrue(greeting.paragraph_format.keep_with_next)
            self.assertEqual(greeting.paragraph_format.space_after.pt, 12)
            self.assertEqual(opening.paragraph_format.space_after.pt, 12)
            self.assertTrue(closing.paragraph_format.keep_with_next)
            self.assertTrue(name.runs[0].bold)
            self.assertEqual(heading.runs[0].font.size.pt, 16)
            self.assertEqual(opening.runs[0].font.size.pt, 12)
            self.assertEqual(opening.runs[0].font.name, "Work Sans")

    def test_resume_uses_reference_typography_geometry_and_category_grid(self):
        content = """NAME: Alex Example
ROLE: Product Designer
CONTACT: alex@example.com

PROFESSIONAL SUMMARY
Designer focused on accessible digital experiences.

---

CORE SKILLS
CATEGORY: Product Design | UX research, prototyping, design systems
CATEGORY: Collaboration | stakeholder workshops, presentations
"""
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "resume.docx"
            _build_resume_docx(content, path)
            document = Document(path)

            name = paragraph_starting_with(document, "ALEX EXAMPLE")
            summary = paragraph_starting_with(document, "Designer focused")
            section = document.sections[0]

            self.assertEqual(name.runs[0].font.size.pt, 13)
            self.assertEqual(str(name.runs[0].font.color.rgb), "205968")
            self.assertEqual(summary.runs[0].font.size.pt, 9)
            self.assertEqual(summary.runs[0].font.name, "Poppins")
            self.assertAlmostEqual(section.left_margin.inches, 0.55, places=2)
            self.assertAlmostEqual(section.right_margin.inches, 0.55, places=2)
            self.assertEqual(len(document.tables), 1)
            self.assertEqual(len(document.tables[0].columns), 2)
            self.assertIn("Product Design", document.tables[0].cell(0, 0).text)
            self.assertIn("Collaboration", document.tables[0].cell(0, 1).text)
            self.assertIn(
                "UX research, prototyping, design systems",
                document.tables[0].cell(0, 0).text,
            )
            self.assertNotIn(";", document.tables[0].cell(0, 0).text)
            self.assertIn("PAGE", document.sections[0].footer._element.xml)

    def test_resume_hyperlinks_bare_website_without_linking_email_domain(self):
        content = """NAME: Alex Example
ROLE: Product Designer
CONTACT: alex@example.com
LINKS: alex.example
"""
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "resume.docx"
            _build_resume_docx(content, path)
            document = Document(path)
            hyperlink_targets = {
                relationship.target_ref
                for relationship in document.part.rels.values()
                if relationship.reltype == RT.HYPERLINK
            }

            self.assertIn("https://alex.example/", hyperlink_targets)
            self.assertNotIn("https://example.com", hyperlink_targets)


if __name__ == "__main__":
    unittest.main()
