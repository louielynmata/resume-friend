import tempfile
import unittest
from pathlib import Path
from unittest.mock import AsyncMock, patch

import pymupdf
from docx import Document

from backend.config import settings
from backend.qa_models import QASeverity, VisualQAResult
from backend.services.artifact_qa_service import (
    inspect_artifacts,
    inspect_artifacts_visually,
)


def write_docx(path: Path, label: str) -> None:
    document = Document()
    document.add_heading(label, level=1)
    document.add_paragraph(
        "This artifact contains enough readable text to exercise the document "
        "validator and confirm the generated Word package can be reopened."
    )
    document.save(path)


def write_pdf(path: Path, pages: int) -> None:
    with pymupdf.open() as document:
        for index in range(pages):
            page = document.new_page()
            page.insert_text(
                (72, 72),
                f"Page {index + 1} contains readable validation text for Resume Friend.",
            )
        document.save(path)


class ArtifactQAServiceTests(unittest.TestCase):
    def test_artifacts_pass_with_expected_page_counts(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            resume_docx = root / "resume.docx"
            cover_docx = root / "cover.docx"
            resume_pdf = root / "resume.pdf"
            cover_pdf = root / "cover.pdf"
            write_docx(resume_docx, "Resume")
            write_docx(cover_docx, "Cover Letter")
            write_pdf(resume_pdf, 2)
            write_pdf(cover_pdf, 1)

            result = inspect_artifacts(
                {
                    "resume_docx": str(resume_docx),
                    "cover_letter_docx": str(cover_docx),
                    "resume_pdf": str(resume_pdf),
                    "cover_letter_pdf": str(cover_pdf),
                }
            )

            self.assertEqual(result.resume_pages, 2)
            self.assertEqual(result.cover_letter_pages, 1)
            self.assertFalse(result.blocking_issues)

    def test_resume_page_limit_is_blocking(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            resume_docx = root / "resume.docx"
            cover_docx = root / "cover.docx"
            resume_pdf = root / "resume.pdf"
            cover_pdf = root / "cover.pdf"
            write_docx(resume_docx, "Resume")
            write_docx(cover_docx, "Cover Letter")
            write_pdf(resume_pdf, 3)
            write_pdf(cover_pdf, 1)

            result = inspect_artifacts(
                {
                    "resume_docx": str(resume_docx),
                    "cover_letter_docx": str(cover_docx),
                    "resume_pdf": str(resume_pdf),
                    "cover_letter_pdf": str(cover_pdf),
                }
            )

            page_limit = [
                issue for issue in result.issues if issue.code == "PDF_PAGE_LIMIT"
            ]
            self.assertEqual(len(page_limit), 1)
            self.assertEqual(page_limit[0].severity, QASeverity.ERROR)


class VisualArtifactQAServiceTests(unittest.IsolatedAsyncioTestCase):
    async def test_visual_qa_renders_pdf_pages_for_the_reviewer(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            resume_pdf = root / "resume.pdf"
            write_pdf(resume_pdf, 2)
            original_temp_dir = settings.qa_temp_dir
            original_reference_dir = settings.reference_dir
            settings.qa_temp_dir = str(root / "rendered")
            settings.reference_dir = str(root / "references")

            async def verify_images(*args, **kwargs):
                image_paths = kwargs["image_paths"]
                self.assertEqual(len(image_paths), 2)
                self.assertTrue(all(path.exists() for path in image_paths))
                return VisualQAResult(passed=True, summary="Looks good")

            try:
                with patch(
                    "backend.services.artifact_qa_service.generate_structured",
                    new=AsyncMock(side_effect=verify_images),
                ):
                    result = await inspect_artifacts_visually(
                        provider="ollama",
                        docs={"resume_pdf": str(resume_pdf)},
                    )
            finally:
                settings.qa_temp_dir = original_temp_dir
                settings.reference_dir = original_reference_dir

            self.assertTrue(result.passed)

    async def test_visual_qa_includes_reference_pages_and_labels(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            resume_pdf = root / "resume.pdf"
            references = root / "references"
            references.mkdir()
            reference_pdf = references / "Design Resume Reference.pdf"
            write_pdf(resume_pdf, 1)
            write_pdf(reference_pdf, 2)
            original_temp_dir = settings.qa_temp_dir
            original_reference_dir = settings.reference_dir
            settings.qa_temp_dir = str(root / "rendered")
            settings.reference_dir = str(references)

            async def verify_reference_context(*args, **kwargs):
                self.assertEqual(len(kwargs["image_paths"]), 3)
                self.assertIn("Reference: Design Resume Reference", args[2])
                return VisualQAResult(passed=True, summary="Matches reference")

            try:
                with patch(
                    "backend.services.artifact_qa_service.generate_structured",
                    new=AsyncMock(side_effect=verify_reference_context),
                ):
                    result = await inspect_artifacts_visually(
                        provider="ollama",
                        docs={"resume_pdf": str(resume_pdf)},
                    )
            finally:
                settings.qa_temp_dir = original_temp_dir
                settings.reference_dir = original_reference_dir

            self.assertTrue(result.passed)


if __name__ == "__main__":
    unittest.main()
