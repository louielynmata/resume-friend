import re
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from .document_normalization import normalize_resume_bullets


_RESUME_FONT = "Poppins"
_RESUME_BODY_SIZE = 9
_RESUME_ACCENT = RGBColor(0x20, 0x59, 0x68)
_COVER_FONT = "Work Sans"
_COVER_BODY_SIZE = 12


async def build_documents(
    ai_response: str,
    owner_name: str,
    position_slug: str,
    output_dir: Path,
) -> dict:
    resume_text = _extract_tag(ai_response, "RESUME")
    cl_text = _extract_tag(ai_response, "COVER_LETTER")

    result: dict = {}

    if resume_text:
        docx_path = _resolve_output_path(output_dir / f"{owner_name}_Resume_{position_slug}.docx")
        docx_path = _build_resume_docx(resume_text, docx_path)
        result["resume_docx"] = str(docx_path)
        result["resume_pdf"] = _to_pdf(docx_path)
    else:
        raise ValueError("AI response did not contain a <RESUME> section.")

    if cl_text:
        cl_path = _resolve_output_path(output_dir / f"{owner_name}_CoverLetter_{position_slug}.docx")
        cl_path = _build_cover_letter_docx(cl_text, cl_path)
        result["cover_letter_docx"] = str(cl_path)
        result["cover_letter_pdf"] = _to_pdf(cl_path)
    else:
        raise ValueError("AI response did not contain a <COVER_LETTER> section.")

    return result


# ── Parsing helpers ────────────────────────────────────────────────────────────

def _extract_tag(text: str, tag: str) -> str:
    match = re.search(rf"<{tag}>(.*?)</{tag}>", text, re.DOTALL)
    return match.group(1).strip() if match else ""


def _resolve_output_path(path: Path) -> Path:
    if not path.exists():
        return path
    if _is_locked_word_document(path):
        return _next_available_path(path)
    return path


def _next_available_path(path: Path) -> Path:
    for index in range(2, 1000):
        candidate = path.with_name(f"{path.stem}_{index}{path.suffix}")
        if not candidate.exists():
            return candidate
    raise OSError(f"Could not find an available output filename for {path.name}.")


def _is_locked_word_document(path: Path) -> bool:
    names = [f"~${path.name}"]
    if len(path.name) >= 2:
        names.append(f"~${path.name[2:]}")
    return any(path.with_name(name).exists() for name in names)


# ── Text normalisation ─────────────────────────────────────────────────────────

def _normalize_document_text(content: str) -> list[str]:
    normalized, _ = normalize_resume_bullets(content)
    normalized = normalized.strip()
    # Replace em dash and en dash used as sentence separators with a plain hyphen.
    # En dash between digits is kept (date ranges like 2024-2026 are fine as hyphens).
    normalized = re.sub(r"\s*—\s*", " - ", normalized)         # em dash (U+2014)
    normalized = re.sub(r"(?<!\d)\s*–\s*(?!\d)", " - ", normalized)  # en dash not between digits
    normalized = re.sub(r"^```(?:[a-zA-Z0-9_-]+)?\s*", "", normalized)
    normalized = re.sub(r"\s*```$", "", normalized)

    # Split "---SECTION" combined lines into separate divider + section lines
    expanded: list[str] = []
    for raw in normalized.split("\n"):
        m = re.match(r"^(-{3,})\s*([^-\s].+)$", raw.strip())
        if m:
            expanded.append(m.group(1))
            expanded.append(m.group(2).strip())
        else:
            expanded.append(raw)

    lines: list[str] = []
    for raw_line in expanded:
        line = raw_line.strip()
        if _ARTIFACT_PAT.match(line):  # drop system-prompt separator echoes
            continue
        if _META_COMMENTARY_PAT.search(line):  # drop AI self-commentary embedded in body text
            continue
        # Bullet markers were canonicalized before line-level rendering.
        # Unwrap markdown links
        line = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", line)
        # Strip markdown formatting underscores/backticks/tildes — but KEEP asterisks for inline bold.
        # Only strip underscores at word boundaries (markdown _italic_), not inside URLs/words.
        line = re.sub(r"(?<![a-zA-Z0-9])_+(?![a-zA-Z0-9])|[`~]+", "", line)
        line = re.sub(r"\s{2,}", " ", line).strip()
        lines.append(line)

    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()

    return lines


def _strip_bold(text: str) -> str:
    """Remove ** markers for classification purposes only."""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text).replace("*", "").strip()


# ── Resume line classifier ─────────────────────────────────────────────────────

_DATE_PAT = re.compile(
    r"^(Jan(uary)?|Feb(ruary)?|Mar(ch)?|Apr(il)?|May|Jun(e)?|Jul(y)?|"
    r"Aug(ust)?|Sep(tember)?|Oct(ober)?|Nov(ember)?|Dec(ember)?|"
    r"Summer|Fall|Winter|Spring|\d{4})",
    re.IGNORECASE,
)

_CONTACT_PREFIX = re.compile(
    r"^(email|phone|tel|mobile|location|city|address|"
    r"linkedin|github|gitlab|portfolio|website|url|links?)\s*[:\-]\s*",
    re.IGNORECASE,
)

_CONTACT_CONTENT = re.compile(
    r"linkedin\.com|github\.com|gitlab\.com|@\S+\.\S+|\+\d[\d\s\-\.]{6,}",
    re.IGNORECASE,
)

_ALLOWED_LOWER = {"and", "or", "of", "the", "by", "in", "with", "to", "a", "&"}

# System-prompt section labels that the AI occasionally echoes — never treat as headers
_BLOCKED_SECTION_NAMES = frozenset({
    "INSTRUCTIONS", "WRITING STYLE EXAMPLES", "WRITING STYLE",
    "TRANSCRIPT", "ANALYSIS",
})

# Lines matching this pattern are system-prompt artifacts or AI meta-commentary — drop them entirely
_ARTIFACT_PAT = re.compile(
    r"^(instructions|writing\s+style\s+examples?|transcript|analysis)\s*[-─═]*$"
    r"|^\(this section is reserved"
    r"|^\(note[:\s]"
    r"|^\[note[:\s]",
    re.IGNORECASE,
)

# Phrases that indicate AI self-commentary embedded in resume body text — drop any line containing these
_META_COMMENTARY_PAT = re.compile(
    r"placeholder\s+section"
    r"|reserved\s+to\s+structure"
    r"|left\s+to\s+structure"
    r"|no\s+dates\s+or\s+details\s+are\s+added"
    r"|ensuring\s+all\s+roles\s+appear"
    r"|as\s+per\s+(?:formatting|instructions?)\s+rules"
    r"|this\s+section\s+is\s+(?:left|reserved|used|here)",
    re.IGNORECASE,
)


def _is_section_header(clean: str) -> bool:
    """ALL CAPS section header — allows lowercase connectors like 'and'."""
    if len(clean) < 3 or "|" in clean or clean.startswith("●"):
        return False
    if "(" in clean:  # institution/company names have parentheses; section headers don't
        return False
    letters = re.sub(r"[^A-Za-z\s]", "", clean).strip()
    if not letters:
        return False
    if letters.upper() in _BLOCKED_SECTION_NAMES:  # reject system-prompt echo artifacts
        return False
    words = letters.split()
    if len(words) > 5:  # section headers are short; long phrases are company/institution names
        return False
    return (
        all(w.isupper() or w.lower() in _ALLOWED_LOWER for w in words)
        and any(w.isupper() for w in words)
        and not _DATE_PAT.match(clean)
    )


def _resume_line_kind(line: str, index: int) -> tuple[str, str]:
    if not line:
        return "blank", ""

    # Work on a ** -stripped version for pattern matching; keep original for rendering
    clean = _strip_bold(line)

    # ── Horizontal divider (---, ─── etc.) ────────────────────────────────────
    if re.fullmatch(r"[-─═]{3,}", clean):
        return "divider", ""

    # ── Markdown heading shortcuts ─────────────────────────────────────────────
    if re.match(r"^#\s+", clean):
        return "name", re.sub(r"^#+\s*", "", clean)

    if re.match(r"^##\s+", clean):
        return "section", re.sub(r"^#+\s*", "", clean)

    if re.match(r"^###\s+", clean):
        return "entry", re.sub(r"^#+\s*", "", clean)

    # ── Structured header markers ──────────────────────────────────────────────
    m = re.match(r"^\[?NAME\s*:\s*(.+?)\]?$", clean, re.IGNORECASE)
    if m:
        return "name", m.group(1).strip()

    m = re.match(r"^\[?(?:FULL NAME)\s*:\s*(.+?)\]?$", clean, re.IGNORECASE)
    if m:
        return "name", m.group(1).strip()

    m = re.match(r"^\[?ROLE\s*:\s*(.+?)\]?$", clean, re.IGNORECASE)
    if m:
        return "role", m.group(1).strip()

    m = re.match(r"^\[?TAGLINE\s*:\s*(.+?)\]?$", clean, re.IGNORECASE)
    if m:
        return "tagline", m.group(1).strip()

    m = re.match(r"^\[?(?:CONTACT|CONTACT INFO)\s*:\s*(.+?)\]?$", clean, re.IGNORECASE)
    if m:
        return "contact", m.group(1).strip()

    m = re.match(r"^\[?(?:LINKS?|PORTFOLIO|DESIGNER PORTFOLIO(?:\s+LINK)?)\s*:\s*(.+?)\]?$", clean, re.IGNORECASE)
    if m:
        # Keep the full line for DESIGNER PORTFOLIO so the label stays visible in the rendered doc
        if re.match(r"^\[?DESIGNER PORTFOLIO", clean, re.IGNORECASE):
            return "contact", clean
        return "contact", m.group(1).strip()

    # Compact skills/toolkit groups. The builder renders consecutive CATEGORY
    # lines as a reference-style column grid without asking the model to emit a
    # Markdown or ASCII table.
    category_m = re.match(r"^CATEGORY\s*:\s*(.+?)\s*\|\s*(.+)$", line, re.IGNORECASE)
    if category_m:
        return "category", f"{category_m.group(1).strip()} | {category_m.group(2).strip()}"

    # ── Contact prefix labels (Email:, Phone:, Location:, LinkedIn: etc.) ──────
    if _CONTACT_PREFIX.match(clean):
        val = _CONTACT_PREFIX.sub("", clean).strip()
        return "contact", val

    # ── Lines that look like contact info within the header block (first 15) ──
    if index < 15 and _CONTACT_CONTENT.search(clean):
        return "contact", clean

    # ── Section headers ────────────────────────────────────────────────────────
    if _is_section_header(clean):
        return "section", clean

    # ── Bullet points ──────────────────────────────────────────────────────────
    bullet_m = re.match(r"^[●•]\s+(.*)", line)
    if bullet_m:
        return "bullet", bullet_m.group(1).strip()

    if line.startswith("- "):
        return "bullet", line[2:].strip()

    # ── Early contact lines (pipe-separated, within first 6 lines) ────────────
    if index <= 5 and "|" in clean:
        return "contact", clean

    # ── Entry headers (pipe-separated, after header block) ────────────────────
    if "|" in clean and index > 5:
        first_segment = clean.split("|")[0].strip()
        if _DATE_PAT.match(first_segment):
            return "date_range", clean
        return "entry", line

    # ── Role-with-date lines: "ROLE TITLE - dates" (Pattern B, after header) ──
    # Must come after section-header check; those fail because mixed-case dates
    # break the all-caps requirement.
    if index > 5:
        role_m = re.match(r"^([A-Z][A-Z\s/&,]+?)\s*[-–—]\s*(.+)$", clean)
        if role_m:
            role_words = role_m.group(1).strip().split()
            if role_words and all(w.isupper() for w in role_words if w.isalpha()):
                return "role_line", line

    # ── Standalone date-range lines (after header block) ──────────────────────
    if index > 5 and _DATE_PAT.match(clean):
        return "date_range", clean

    return "body", line


# ── Hyperlink helpers ──────────────────────────────────────────────────────────

_URL_LINK_PAT = re.compile(
    r"(https?://[^\s|]+|(?:linkedin|github|gitlab)\.com/[^\s|]+)",
    re.IGNORECASE,
)

# Maps specific URLs to short display labels. Portfolio folder URL is intentionally absent
# so the full URL is displayed as-is (user preference: show the real link).
_URL_LABELS: dict[str, str] = {}


def _ensure_url(text: str) -> str:
    return text if text.startswith("http") else f"https://{text}"


def _add_hyperlink(
    paragraph,
    display: str,
    url: str,
    size: int | float,
    *,
    font_family: str = _RESUME_FONT,
) -> None:
    """Insert a clickable hyperlink run into an existing paragraph."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_family)
    rFonts.set(qn("w:hAnsi"), font_family)
    rFonts.set(qn("w:eastAsia"), font_family)
    rFonts.set(qn("w:cs"), font_family)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    rPr.append(sz)

    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = display
    t.set(qn("xml:space"), "preserve")
    r.append(t)

    hl.append(r)
    paragraph._p.append(hl)


def _add_contact_line(
    paragraph,
    text: str,
    size: int | float,
    *,
    font_family: str = _RESUME_FONT,
) -> None:
    """Render a contact/links line, converting URL-shaped tokens to hyperlinks."""
    segments = _URL_LINK_PAT.split(text)
    for seg in segments:
        if not seg:
            continue
        if _URL_LINK_PAT.fullmatch(seg):
            url = _ensure_url(seg)
            display = _URL_LABELS.get(url, seg)
            _add_hyperlink(
                paragraph,
                display,
                url,
                size,
                font_family=font_family,
            )
        else:
            run = paragraph.add_run(seg)
            _set_run_font(run, size=size, font_family=font_family)


# ── Inline bold renderer ───────────────────────────────────────────────────────

def _add_inline_runs(
    paragraph,
    text: str,
    size: int | float,
    bold_base: bool = False,
    color: Optional[RGBColor] = None,
    *,
    font_family: str = _RESUME_FONT,
) -> None:
    """Split text on **...** and add bold/normal runs to an existing paragraph."""
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = bold_base or (i % 2 == 1)
        _set_run_font(run, size=size, font_family=font_family)
        if color and not (i % 2 == 1):
            run.font.color.rgb = color


# ── Resume .docx ───────────────────────────────────────────────────────────────

def _next_rendered_kind(lines: list[str], index: int) -> str | None:
    for next_index in range(index + 1, len(lines)):
        next_line = lines[next_index].strip()
        kind, _ = _resume_line_kind(next_line, next_index)
        if kind != "blank":
            return kind
    return None


def _build_resume_docx(content: str, path: Path) -> Path:
    doc = Document()
    # Reference PDFs use a compact Poppins layout with approximately 0.5-0.75
    # inch margins and 8.5-9 pt body copy.
    _set_margins(doc, top=0.55, bottom=0.5, left=0.55, right=0.55)
    _set_default_font(doc, font_family=_RESUME_FONT, size=_RESUME_BODY_SIZE)
    _configure_resume_bullet_style(doc)
    _add_page_number_footer(doc)

    lines = _normalize_document_text(content)
    i = 0
    current_section = ""
    while i < len(lines):
        line = lines[i].strip()
        kind, value = _resume_line_kind(line, i)
        next_kind = _next_rendered_kind(lines, i)

        if kind == "blank":
            i += 1
            continue

        if kind == "name":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value.upper())
            run.bold = True
            _set_run_font(run, size=13)
            run.font.color.rgb = _RESUME_ACCENT
            _set_para_spacing(p, before=0, after=0)
            # Hard rule directly under name — the signature style element
            _add_full_rule(doc, before=1, after=3)

        elif kind == "role":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            run.bold = True
            _set_run_font(run, size=9)
            run.font.color.rgb = _RESUME_ACCENT
            _set_para_spacing(p, before=0, after=1)

        elif kind == "tagline":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            _set_run_font(run, size=9)
            _set_para_spacing(p, before=0, after=1)

        elif kind == "contact":
            # Strip any remaining label prefixes that sneak through
            display = _CONTACT_PREFIX.sub("", value).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_contact_line(p, display, size=8.5)
            _set_para_spacing(p, before=0, after=1)

        elif kind == "divider":
            _add_full_rule(doc, before=4, after=4)

        elif kind == "section":
            current_section = value.upper()
            _add_section_header(doc, value)

        elif kind == "category":
            categories: list[tuple[str, str]] = []
            while i < len(lines):
                candidate_kind, candidate_value = _resume_line_kind(
                    lines[i].strip(),
                    i,
                )
                if candidate_kind != "category":
                    break
                label, details = candidate_value.split("|", 1)
                categories.append((label.strip(), details.strip()))
                i += 1
            columns = 3 if current_section == "TOOLKIT" else 2
            _add_category_grid(doc, categories, columns=columns)
            continue

        elif kind == "entry":
            _add_entry_header(
                doc,
                value,
                keep_with_next=next_kind in {
                    "entry",
                    "role_line",
                    "date_range",
                    "bullet",
                    "body",
                },
            )

        elif kind == "role_line":
            _add_role_with_date(doc, value)

        elif kind == "date_range":
            p = doc.add_paragraph()
            run = p.add_run(_strip_bold(value))
            _set_run_font(run, size=_RESUME_BODY_SIZE)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            _set_para_spacing(p, before=0, after=1)
            _set_pagination(
                p,
                keep_with_next=next_kind in {"bullet", "body"},
                keep_together=True,
            )

        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _set_para_spacing(p, before=0, after=0)
            _set_pagination(p, keep_together=True)
            _add_inline_runs(p, value, size=_RESUME_BODY_SIZE)

        else:  # body
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_spacing(p, before=0, after=1)
            _set_pagination(p, keep_together=True)
            _add_inline_runs(p, value, size=_RESUME_BODY_SIZE)

        i += 1

    return _save_document(doc, path)


def _add_section_header(doc: Document, text: str) -> None:
    text = re.sub(r"^-+\s*", "", text)   # strip leading dashes
    text = re.sub(r"\s*-+$", "", text)   # strip trailing dashes
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph()
    _set_para_spacing(p, before=4, after=1)
    _set_pagination(p, keep_with_next=True, keep_together=True)
    run = p.add_run(text.upper())
    run.bold = True
    _set_run_font(run, size=9)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _add_role_with_date(doc: Document, text: str) -> None:
    """Render 'ROLE TITLE - dates' as bold role + gray dates on one line."""
    clean = _strip_bold(text)
    m = re.match(r"^([A-Z][A-Z\s/&,]+?)\s*[-–—]\s*(.+)$", clean)
    p = doc.add_paragraph()
    _set_para_spacing(p, before=1, after=0)
    _set_pagination(p, keep_with_next=True, keep_together=True)
    if m:
        run = p.add_run(m.group(1).strip())
        run.bold = True
        _set_run_font(run, size=_RESUME_BODY_SIZE)
        run2 = p.add_run("  -  " + m.group(2).strip())
        _set_run_font(run2, size=_RESUME_BODY_SIZE)
        run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    else:
        _add_inline_runs(p, text, size=_RESUME_BODY_SIZE)


def _add_entry_header(
    doc: Document,
    text: str,
    *,
    keep_with_next: bool = False,
) -> None:
    clean = _strip_bold(text)
    parts = [p.strip() for p in clean.split("|")]
    p = doc.add_paragraph()
    _set_para_spacing(p, before=4, after=0)
    _set_pagination(
        p,
        keep_with_next=keep_with_next,
        keep_together=True,
    )
    if parts:
        run = p.add_run(parts[0])
        run.bold = True
        _set_run_font(run, size=_RESUME_BODY_SIZE)
    if len(parts) > 1:
        run = p.add_run("  |  " + "  |  ".join(parts[1:]))
        _set_run_font(run, size=_RESUME_BODY_SIZE)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _add_category_grid(
    doc: Document,
    categories: list[tuple[str, str]],
    *,
    columns: int,
) -> None:
    if not categories:
        return

    columns = max(1, min(columns, len(categories)))
    rows = (len(categories) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    usable_width = 8.5 - 1.1
    column_width = usable_width / columns
    _set_table_borders_none(table)

    for index, cell in enumerate(cell for row in table.rows for cell in row.cells):
        cell.width = Inches(column_width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_cell_margins(
            cell,
            top=0,
            bottom=80,
            left=0 if index % columns == 0 else 100,
            right=140 if index % columns < columns - 1 else 0,
        )
        paragraph = cell.paragraphs[0]
        _set_para_spacing(paragraph, before=0, after=0)
        _set_pagination(paragraph, keep_together=True)

        if index >= len(categories):
            continue

        label, details = categories[index]
        label_run = paragraph.add_run(label)
        label_run.bold = True
        _set_run_font(label_run, size=9)
        label_run.add_break()
        _add_inline_runs(paragraph, details, size=8.5)

    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)


def _set_table_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        borders.append(border)


def _set_cell_margins(
    cell,
    *,
    top: int,
    bottom: int,
    left: int,
    right: int,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _add_full_rule(doc: Document, before: int = 0, after: int = 0) -> None:
    """Standalone full-width horizontal rule paragraph."""
    p = doc.add_paragraph()
    _set_para_spacing(p, before=before, after=after)
    # Collapse paragraph body to 1 pt so the border renders as a thin line,
    # not a thick block (empty paragraph inherits the Normal 10 pt line height).
    p.paragraph_format.line_spacing = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Cover Letter .docx ─────────────────────────────────────────────────────────

def _cover_letter_line_kind(line: str) -> str:
    if re.match(r"^cover letter$", line, re.IGNORECASE):
        return "heading"
    if re.match(r"^(?:to|dear)\b", line, re.IGNORECASE):
        return "greeting"
    if re.match(r"^cheers", line, re.IGNORECASE):
        return "signoff"
    if re.match(r"^sincerely", line, re.IGNORECASE):
        return "closing"
    return "body"


def _cover_letter_lines(content: str) -> list[str]:
    raw = content.replace("\r\n", "\n").replace("\r", "\n").strip()
    raw = re.sub(r"\s*—\s*", " - ", raw)                         # em dash (U+2014)
    raw = re.sub(r"(?<!\d)\s*–\s*(?!\d)", " - ", raw)            # en dash not between digits
    raw = re.sub(r"^```(?:[a-zA-Z0-9_-]+)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    cleaned = []
    for raw_line in raw.split("\n"):
        line = re.sub(r"^#{1,6}\s*", "", raw_line).strip()
        line = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", line)
        line = re.sub(r"(?<![a-zA-Z0-9])_+(?![a-zA-Z0-9])|[`~]+", "", line)   # strip markdown formatting, preserve in-word/URL underscores
        line = re.sub(r"\s{2,}", " ", line)
        cleaned.append(line)
    return cleaned


def _build_cover_letter_docx(content: str, path: Path) -> Path:
    doc = Document()
    # Match the retained cover-letter reference: Work Sans, readable 12 pt
    # body copy, a 16 pt heading, and one-inch side margins.
    _set_margins(doc, top=1.15, bottom=0.75, left=1.0, right=1.0)
    _set_default_font(doc, font_family=_COVER_FONT, size=_COVER_BODY_SIZE)

    lines = _cover_letter_lines(content)
    in_closing = False

    for line in lines:
        if not line:
            continue

        kind = _cover_letter_line_kind(line)
        clean = _strip_bold(line)

        if kind == "heading":
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.bold = True
            _set_run_font(run, size=16, font_family=_COVER_FONT)
            _set_para_spacing(p, before=0, after=20)
            _set_pagination(p, keep_with_next=True, keep_together=True)

        elif kind == "greeting":
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.bold = True
            _set_run_font(run, size=_COVER_BODY_SIZE, font_family=_COVER_FONT)
            _set_para_spacing(p, before=0, after=12)
            _set_pagination(p, keep_with_next=True, keep_together=True)

        elif kind == "signoff":
            p = doc.add_paragraph()
            _set_run_font(
                p.add_run(clean),
                size=_COVER_BODY_SIZE,
                font_family=_COVER_FONT,
            )
            _set_para_spacing(p, before=2, after=8)
            _set_pagination(p, keep_with_next=True, keep_together=True)
            in_closing = True

        elif kind == "closing":
            p = doc.add_paragraph()
            _set_run_font(
                p.add_run(clean),
                size=_COVER_BODY_SIZE,
                font_family=_COVER_FONT,
            )
            _set_para_spacing(p, before=0, after=0)
            _set_pagination(p, keep_with_next=True, keep_together=True)
            in_closing = True

        else:
            if in_closing:
                # Closing block: bold the name line, plain for email/url
                is_url_or_email = line.startswith("http") or re.match(r".+@.+\..+", line)
                p = doc.add_paragraph()
                if is_url_or_email and line.startswith("http"):
                    _add_hyperlink(
                        p,
                        clean,
                        clean,
                        size=_COVER_BODY_SIZE,
                        font_family=_COVER_FONT,
                    )
                elif is_url_or_email:
                    _add_hyperlink(
                        p,
                        clean,
                        f"mailto:{clean}",
                        size=_COVER_BODY_SIZE,
                        font_family=_COVER_FONT,
                    )
                else:
                    run = p.add_run(clean)
                    run.bold = True
                    _set_run_font(
                        run,
                        size=_COVER_BODY_SIZE,
                        font_family=_COVER_FONT,
                    )
                _set_para_spacing(p, before=0, after=0)
                _set_pagination(p, keep_together=True)
            else:
                # Body paragraph — render inline bold from **...**
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.08
                _set_para_spacing(p, before=0, after=12)
                _set_pagination(p, keep_together=True)
                _add_inline_runs(
                    p,
                    line,
                    size=_COVER_BODY_SIZE,
                    font_family=_COVER_FONT,
                )

    return _save_document(doc, path)


# ── PDF conversion ─────────────────────────────────────────────────────────────

def _to_pdf(docx_path: Path) -> Optional[str]:
    try:
        from docx2pdf import convert
        pdf_path = docx_path.with_suffix(".pdf")
        convert(str(docx_path), str(pdf_path))
        return str(pdf_path) if pdf_path.exists() else None
    except Exception:
        return None


# ── Helpers ────────────────────────────────────────────────────────────────────

def _set_margins(doc: Document, top: float, bottom: float, left: float, right: float) -> None:
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def _set_default_font(
    doc: Document,
    *,
    font_family: str,
    size: int | float,
) -> None:
    normal_style = doc.styles["Normal"]
    _apply_font_family(normal_style.font, font_family)
    normal_style.font.size = Pt(size)
    normal_style.paragraph_format.line_spacing = 1.0


def _set_run_font(
    run,
    size: int | float,
    *,
    font_family: str = _RESUME_FONT,
) -> None:
    _apply_font_family(run.font, font_family)
    run.font.size = Pt(size)


def _apply_font_family(font, font_family: str) -> None:
    font.name = font_family
    font._element.rPr.rFonts.set(qn("w:ascii"), font_family)
    font._element.rPr.rFonts.set(qn("w:hAnsi"), font_family)
    font._element.rPr.rFonts.set(qn("w:eastAsia"), font_family)
    font._element.rPr.rFonts.set(qn("w:cs"), font_family)


def _configure_resume_bullet_style(doc: Document) -> None:
    style = doc.styles["List Bullet"]
    _apply_font_family(style.font, _RESUME_FONT)
    style.font.size = Pt(_RESUME_BODY_SIZE)
    style.paragraph_format.left_indent = Inches(0.28)
    style.paragraph_format.first_line_indent = Inches(-0.18)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0


def _add_page_number_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(paragraph, before=0, after=0)

        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")

        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), _RESUME_FONT)
        fonts.set(qn("w:hAnsi"), _RESUME_FONT)
        run_properties.append(fonts)

        size = OxmlElement("w:sz")
        size.set(qn("w:val"), "14")
        run_properties.append(size)

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "CCCCCC")
        run_properties.append(color)

        run.append(run_properties)
        text = OxmlElement("w:t")
        text.text = "1"
        run.append(text)
        field.append(run)
        paragraph._p.append(field)


def _save_document(doc: Document, path: Path) -> Path:
    try:
        doc.save(str(path))
        return path
    except PermissionError:
        fallback_path = _next_available_path(path)
        doc.save(str(fallback_path))
        return fallback_path


def _set_para_spacing(p, before: int = 0, after: int = 4) -> None:
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


def _set_pagination(
    paragraph,
    *,
    keep_with_next: bool = False,
    keep_together: bool = False,
) -> None:
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = keep_together
    paragraph.paragraph_format.widow_control = True
